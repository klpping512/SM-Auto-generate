"""Dump the seller-topic source documents so the real user-side topics can be
read without guessing.  Read-only; writes nothing back to the source folder.
"""
from __future__ import annotations

import sys
from pathlib import Path

FOLDER = Path("/Users/ylanlll/Desktop/十个主题优化版")


def dump_docx(path: Path) -> None:
    import docx

    document = docx.Document(str(path))
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                print(" | ".join(cells))


def dump_xlsx(path: Path) -> None:
    import openpyxl

    book = openpyxl.load_workbook(str(path), data_only=True)
    for sheet in book.worksheets:
        print(f"--- sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if cells:
                print(" | ".join(cells))


def main() -> None:
    targets = sys.argv[1:] or [item.name for item in sorted(FOLDER.iterdir()) if not item.name.startswith("~$")]
    for name in targets:
        path = FOLDER / name
        print(f"\n{'=' * 70}\n== {name}\n{'=' * 70}")
        try:
            if path.suffix.lower() == ".docx":
                dump_docx(path)
            elif path.suffix.lower() == ".xlsx":
                dump_xlsx(path)
            else:
                print(f"[跳过：{path.suffix} 需要先另存为 .docx/.xlsx]")
        except Exception as exc:
            print(f"[读取失败：{type(exc).__name__}: {exc}]")


if __name__ == "__main__":
    main()
